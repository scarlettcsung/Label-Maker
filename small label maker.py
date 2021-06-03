# Created by: Scarlet Sung
# Last update: May 25, 2021

print('~INTRODUCTION~')
print('This program will take texts, separated by commas (,)')
print('and print all text onto a small label template.')
print('For example, input \"Hello,World\" will create a label file with two labels, \"Hello\" and \"World\"')
print('Hope this helps!\n')
print('Last updated: May 25, 2021\n')

from docx import Document
import docx2txt

inputted_os = input(str('Copy and Paste all label text here, separated with a comma (,) in between: '))
remain_mast = len(inputted_os)
inputted_os = inputted_os.split(',')

remain = remain_mast

if len(inputted_os) > 119:
    remain = 119
    print('Value ends at ')
    print(inputted_os[119])
    print(' Please start program again after this file is finished to make next file.')
    os_no_list = inputted_os[:119]
else:
    os_no_list = inputted_os

template = Document('Small Template.docx')
template.save('small label ' + inputted_os[0] + '.docx')

no_docs = 1

x = 0
y = 0
n = 0

length_list = len(os_no_list)

while n < 119 and n < length_list:

    os_no = os_no_list[n]
    cell_content = os_no

    repeats = 1
    
    while repeats > 0:

        while y <= 12 and repeats > 0:

            if template.tables[0].cell(x,y).text == '':
                template.tables[0].cell(x,y).text = cell_content
                y += 2
                repeats -= 1
                n += 1
    
            else:
                y += 2
        if repeats == 0:
            break
        
        y = 0
        x += 2

    template.save('small label ' + inputted_os[0] + '.docx')

template.save('small label ' + inputted_os[0] + '.docx')
      
