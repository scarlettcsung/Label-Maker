# Created by: Scarlet Sung
# Last update: May 20, 2021

print('~INTRODUCTION~')
print('This program will take participant information and create label documents')
print('The text template of labels may be created regardless of number of participants or number of each type of sample')
print('For automatic Lab ID continuation, the prefix has to be 2 letters (XX, ex. HA, GA etc.). Otherwise, please manually input Lab ID even if it is a continuation.')
print('However, small label documents may only be created for up to 4 participants, with the default number of samples (10Pl, 6CP, 2 Se, 3 Lab Tubes, 5 Exp Tubes), at one time')
print('Large labels made be made for up to 10 participants at one time, with 2 form labels and 1 bag label')
print('The program will also only create labels and a document for one date')
print('For participant numbers above 4 or 10 (for small and large labels respectively) and number of dates above 1, please run the program again for additional documents')
print('If prompt asks for (Y/N), note that any key other than \'N\' will act as \'Y\'. The input is NOT case-sensitive.')
print('Hope this helps!\n')
print('Last update: May 20, 2021\n')

#date
day_of = input(str('Input the Day of Month (DD): '))

month_of = input(str('Input the Month (MM): '))
year_of = input(str('Input the Year (YYYY): '))

date_of = day_of + '-' + month_of + '-' + year_of

#date confirmation
confirm_date = input(str('The date is %s. Is that correct?(Y/N): ' % date_of))
confirm_date = confirm_date.upper()

while confirm_date == 'N':
    day_of = input(str('Input the Day of Month (DD): '))
    month_of = input(str('Input the Month (MM): '))
    year_of = input(str('Input the Year (YYYY): '))

    date_of = day_of + '-' + month_of + '-' + year_of

    confirm_date = input(str('The date is %s. Is that correct?(Y/N): ' % date_of))
    confirm_date = confirm_date.upper()

#number of participants        
num_participants = input('Input number of participants: ')
# num_participants = input('Input number of participants (Maximum 4 per file): ')

num_participants = int(num_participants)
participant_list = []
participant_order = 1

print('')

part_remain = num_participants

#participant info input
while part_remain > 0 :
    participant_dict = {'Participant name': '','labID':'','HKID':''}
    part_name = input(str('Input participant %d\'s name: ' % participant_order))
    if participant_order != 1:
        labID_cont_yn = input(str('Is participant %d\'s Lab ID a continuation?(Y/N): ' % participant_order))
        labID_cont_yn = labID_cont_yn.upper()
        if labID_cont_yn != 'N':
            split_labID = list(part_labID)
            prefix = split_labID[0:2]
            prefix = [''.join(prefix)]
            number_labID = split_labID[2:]
            number_labID = [''.join(number_labID)]
            num_labID = int(number_labID[0])
            num_labID += 1
            part_labID = prefix[0] + str(num_labID)
        else:
            part_labID = input(str('Input participant %d\'s Lab ID: ' % participant_order))
    else:
        part_labID = input(str('Input participant %d\'s Lab ID: ' % participant_order))
    part_HKID = input(str('Input participant %d\'s HKID: ' %participant_order))
    part_info_yn = input(str('Name: %s Lab ID: %s HKID: %s \nIs the information correct?(Y/N): ' % (part_name, part_labID, part_HKID)))

    part_info_yn = part_info_yn.upper()
    
    if part_info_yn == 'N':
        part_name = input(str('Input participant %d\'s name: ' % participant_order))
        part_labID = input(str('Input participant %d\'s Lab ID: ' % participant_order))
        part_HKID = input(str('Input participant %d\'s HKID: ' %participant_order))
        
    participant_dict['Participant name'] = part_name
    participant_dict['labID'] = part_labID
    participant_dict['HKID'] = part_HKID
    participant_list.append(participant_dict)
    participant_order += 1
    part_remain -= 1


print(participant_list)

part_remain = num_participants
n = 0
research_list = []
elite_list = []
serum_list = []
plasma_list = []
cp_list = []
large_label_list = []
large_label_HKUST_list = []

#printing text for labels

while part_remain > 0:
    print(participant_list[n]['Participant name'].upper() + ' ' + participant_list[n]['labID'])
    print('')
    
    print('~Research blood tubes~')
    research_label = participant_list[n]['Participant name'] + '\n' + participant_list[n]['HKID'] + '\n' + participant_list[n]['labID'] 
    research_list.append(research_label)
    print(research_label)
    print('')
    
    print('~Elite blood tubes~')
    elite_label = participant_list[n]['Participant name'] + '\n' + participant_list[n]['HKID'] + '\n' + 'Lab Test'
    elite_list.append(elite_label)
    print(elite_label)
    print('')
    
    print('~Serum tubes~')
    serum_label = participant_list[n]['labID'] + '\n' + 'Se' + '\n' + date_of
    serum_list.append(serum_label)
    print(serum_label)
    print('')
    
    print('~Plasma tubes~')
    plasma_label = participant_list[n]['labID'] + '\n' + 'Pl' + '\n' + date_of
    plasma_list.append(plasma_label)
    print(plasma_label)
    print('')
    
    print('~Cell pellet tubes~')
    cp_label = participant_list[n]['labID'] + '\n' + 'CP' + '\n' + date_of
    cp_list.append(cp_label)
    print(cp_label)
    print('')

    print('~Large labels~')
    large_label = participant_list[n]['Participant name'] + '\n' + participant_list[n]['HKID'] + '\n' + participant_list[n]['labID'] + '\n' + date_of
    large_label_list.append(large_label)
    large_label_HKUST = participant_list[n]['Participant name'] + '\n' + participant_list[n]['HKID'] + '\n' + date_of + '\n' + 'HKUST'
    large_label_HKUST_list.append(large_label_HKUST)
    print(large_label)
    print('')
    print(large_label_HKUST)
    print('')
    
    n += 1
    part_remain -= 1

#asking if small label document should be made
    
document_yn = input(str('Would you like to have a small label document to be produced for you? (default settings): '))
document_yn = document_yn.upper()

#small label document initiation


if num_participants > 4:
    print('Only one document for the first 4 participants will be made per input. Please restart program for remaining participants after first document creation completion.')
    
part_remain = num_participants

no_docs = 1

x = 0
y = 2
n = 0


#small label document creation
    
if document_yn != 'N':

    print("\nDefault settings:")
    print("10 Plasma, 6 CP, 2 Serum, 3 Lab Test Blood Tubes, 5 Experimental Lab Tubes")

#making small labels

    from docx import Document

    import docx2txt

    while no_docs > 0:
        if part_remain > 4:
            parts_in_doc = participant_list[:4]
        else:
            parts_in_doc = participant_list
        
        document = Document('Small Template.docx')

        part_doc_remain = len(parts_in_doc)

        if len(parts_in_doc) >= 2:
            document.tables[0].cell(0,2).text = parts_in_doc[0]['labID'] + '\n' + parts_in_doc[-1]['labID']
        else:
            document.tables[0].cell(0,2).text = parts_in_doc[0]['labID']

        y = 4
        n = 0

        while part_doc_remain > 0:
            document.tables[0].cell(0,y).text = parts_in_doc[n]['labID'] + '\n' + date_of
            part_doc_remain -= 1
            n += 1
            y += 2

        part_doc_remain = len(parts_in_doc)
        n = 0
        x = 2
        y = 2

        for m in range(part_doc_remain):
            y = 2
            
            document.tables[0].cell(x,y).text = parts_in_doc[n]['labID'] + '\n' + 'Se' + '\n' + date_of
            y += 2
            
            while y <= 12:
                document.tables[0].cell(x,y).text = parts_in_doc[n]['labID'] + '\n' + 'Pl' + '\n' + date_of
                y += 2
            x += 2

            y = 2
            document.tables[0].cell(x,y).text = parts_in_doc[n]['labID'] + '\n' + 'Se' + '\n' + date_of
            y += 2
            while y <= 12:
                document.tables[0].cell(x,y).text = parts_in_doc[n]['labID'] + '\n' + 'Pl' + '\n' + date_of
                y+= 2
            x += 2

            y = 2
            
            document.tables[0].cell(x,y).text = parts_in_doc[n]['labID'] + '\n' + 'CP' + '\n' + date_of
            while y <= 12:
                document.tables[0].cell(x,y).text = parts_in_doc[n]['labID'] + '\n' + 'CP' + '\n' + date_of
                y+= 2
            x += 2
            n += 1

        part_doc_remain = len(parts_in_doc)

        document.save('small label ' + date_of + '.docx')

        x = 2
        y = 0
        n = 0
        
        
        while part_doc_remain > 0: 

            repeats = 3

            while repeats > 0 :
                    
                while y <= 12 and repeats > 0:
                    
                    if document.tables[0].cell(x,y).text == '':
                        document.tables[0].cell(x,y).text = parts_in_doc[n]["Participant name"] + '\n' + parts_in_doc[n]['HKID'] + '\n' + 'Lab Test'
                        y += 2
                        repeats -= 1
                        
                        
                    else:
                        y += 2

                if repeats == 0:
                    break                   

                y = 0
                x += 2

            document.save('small label ' + date_of + '.docx')
            
            repeats = 5
           
            while repeats > 0:
                    
                while y <= 12 and repeats > 0:
                    
                    if document.tables[0].cell(x,y).text == '':
                        document.tables[0].cell(x,y).text = parts_in_doc[n]['Participant name'] + '\n' + parts_in_doc[n]['HKID'] + '\n' + parts_in_doc[n]['labID']
                        y += 2
                        repeats -= 1

                    else:
                        y += 2
                if repeats == 0:
                    break                   

                y = 0
                x += 2
                
            document.save('small label ' + date_of + '.docx')
            
            n += 1
            part_doc_remain -= 1
            
                    
        document.save('small label ' + date_of + '.docx')
        no_docs -= 1

#if large label is to be made
document_l_yn = input(str('\nWould you like to have a large label document to be produced for you? (default settings): '))
document_l_yn = document_l_yn.upper()

if num_participants > 10:
    print('Only one document for the first 10 participants will be made per input. Please restart program for remaining participants after first document creation completion.')
    
part_remain = num_participants

no_docs = 1

x = 0
y = 0
n = 0


#large label document creation
    
if document_l_yn != 'N':

    print("\nDefault settings:")
    print("2 Form labels 1 Bag label")

#making large labels

    from docx import Document

    import docx2txt

    while no_docs > 0:
        if part_remain > 10:
            parts_in_doc = participant_list[:10]
        else:
            parts_in_doc = participant_list
        
        document = Document('Large Template.docx')

        part_doc_remain = len(parts_in_doc)

        n = 0
        x = 0
        y = 0

        for m in range(part_doc_remain):
            y = 0
            
            document.tables[0].cell(x,y).text = parts_in_doc[n]['Participant name'] + '\n' + parts_in_doc[n]['HKID'] + '\n' + parts_in_doc[n]['labID'] + '\n' + date_of
            y += 1
            document.tables[0].cell(x,y).text = parts_in_doc[n]['Participant name'] + '\n' + parts_in_doc[n]['HKID'] + '\n' + parts_in_doc[n]['labID'] + '\n' + date_of
            y += 1
            document.tables[0].cell(x,y).text = parts_in_doc[n]['Participant name'] + '\n' + parts_in_doc[n]['HKID'] + '\n' + date_of + '\n' + 'HKUST'

            x += 1
            n += 1

        document.save('large label ' + date_of + '.docx')

        no_docs -= 1
